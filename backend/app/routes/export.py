"""
ResumeIQ — Export & Sharing Routes (FastAPI)
"""

from datetime import datetime, timedelta, timezone
import secrets
from io import BytesIO
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from docx import Document
from pydantic import BaseModel
from sqlalchemy.orm import Session

from ..database import get_db, User, ResumeDraft, ReviewLink, Resume
from ..middleware.auth import get_current_user

router = APIRouter(prefix="", tags=["Export & Share"])

class ShareRequest(BaseModel):
    expiresInHours: int = 24

class ExportRequest(BaseModel):
    draftId: str

@router.post("/export/pdf")
@router.post("/v1/export/pdf")
def export_pdf(payload: ExportRequest, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    raise HTTPException(status_code=status.HTTP_501_NOT_IMPLEMENTED, detail="PDF export is not available yet. Use DOCX export or the browser print action.")

@router.post("/export/docx")
@router.post("/v1/export/docx")
def export_docx(payload: ExportRequest, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    draft = db.query(ResumeDraft).filter(ResumeDraft.id == payload.draftId, ResumeDraft.userId == user.id).first()
    if not draft:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Draft not found.")

    document = Document()
    document.add_heading(draft.title or "Resume", level=0)
    for section in draft.sections or []:
        document.add_heading(section.get("name") or section.get("title") or "Section", level=1)
        if section.get("content"):
            document.add_paragraph(section["content"])
        for item in section.get("items", []):
            heading = " / ".join(value for value in [item.get("title"), item.get("subtitle"), item.get("date")] if value)
            if heading:
                document.add_paragraph(heading).runs[0].bold = True
            for bullet in item.get("bullets", []):
                text = bullet.get("text") if isinstance(bullet, dict) else str(bullet)
                document.add_paragraph(text, style="List Bullet")
        for bullet in section.get("bullets", []):
            text = bullet.get("text") if isinstance(bullet, dict) else str(bullet)
            document.add_paragraph(text, style="List Bullet")

    output = BytesIO()
    document.save(output)
    output.seek(0)
    filename = "resume.docx"
    return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{filename}"'})

@router.post("/share/resume/{resume_id}", status_code=status.HTTP_201_CREATED)
@router.post("/v1/share/resume/{resume_id}", status_code=status.HTTP_201_CREATED)
def create_share_link(
    resume_id: str,
    payload: ShareRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    token = secrets.token_hex(32)
    expires_at = datetime.now(timezone.utc) + timedelta(hours=payload.expiresInHours)

    resume = db.query(Resume).filter(Resume.id == resume_id, Resume.userId == user.id).first()
    if not resume:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Resume not found.")
    if payload.expiresInHours < 1 or payload.expiresInHours > 720:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Expiry must be between 1 and 720 hours.")

    link = ReviewLink(
        token=token,
        resumeId=resume_id,
        expiresAt=expires_at
    )
    db.add(link)
    db.commit()

    return {
        "status": "success",
        "data": {
            "token": token,
            "expiresAt": expires_at.isoformat()
        }
    }

@router.delete("/share/{token}")
@router.delete("/v1/share/{token}")
def revoke_share_link(token: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    link = db.query(ReviewLink).join(Resume, ReviewLink.resumeId == Resume.id).filter(ReviewLink.token == token, Resume.userId == user.id).first()
    if not link:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Share link not found.")
    db.delete(link)
    db.commit()
    return {"status": "success", "message": "Share link revoked."}
