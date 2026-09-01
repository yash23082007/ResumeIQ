"""
ResumeIQ — Unified Resume Document Parser
Parses PDF, DOCX, and TXT files into raw text and structured metadata.
Validates file magic bytes and extracts layout signals.
"""

import os
import io
import re
from typing import Dict, Any, Tuple
from pypdf import PdfReader
import docx

from .section_extractor import extract_sections

def validate_file_signature(content_bytes: bytes, ext: str):
    """Validate file magic bytes to avoid extension spoofing."""
    if ext == ".pdf":
        if not content_bytes.startswith(b"%PDF-"):
            raise ValueError("Invalid PDF file signature. The file appears to be corrupted or renamed.")
    elif ext == ".docx":
        if len(content_bytes) < 4 or content_bytes[:2] != b"PK":
            raise ValueError("Invalid DOCX file signature. Please upload a valid Word (.docx) document.")
    elif ext == ".doc":
        raise ValueError("Legacy binary .doc files are not supported. Please save your file as .docx or .pdf.")

def detect_header_footer_contact(lines: list) -> bool:
    """Detect if email or phone is solely present in header or footer lines."""
    email_regex = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
    phone_regex = re.compile(r"(\+?\d[\d\s\-().]{7,}\d)")

    header_lines = " ".join(lines[:3])
    footer_lines = " ".join(lines[-3:])

    has_footer = bool(email_regex.search(footer_lines) or phone_regex.search(footer_lines))
    has_header = bool(email_regex.search(header_lines) or phone_regex.search(header_lines))

    return has_footer and not has_header

def parse_pdf_bytes(content_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Extract raw text and layout signals from PDF bytes."""
    reader = PdfReader(io.BytesIO(content_bytes))
    page_count = len(reader.pages)
    text_chunks = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        text_chunks.append(page_text)

    raw_text = "\n".join(text_chunks).strip()
    lines = raw_text.split("\n")
    short_line_ratio = len([l for l in lines if 0 < len(l.strip()) < 30]) / max(len(lines), 1)

    layout_meta = {
        "format": "pdf",
        "pageCount": page_count,
        "hasColumns": short_line_ratio > 0.5,
        "hasImages": len(raw_text) < 100 and page_count > 0,
        "contactInHeaderFooter": detect_header_footer_contact(lines),
        "hasNonStandardFonts": False,
        "hasMultiColumnTables": "\t\t" in raw_text or "   |   " in raw_text,
    }

    return raw_text, layout_meta

def parse_docx_bytes(content_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Extract raw text and layout signals from DOCX bytes."""
    doc = docx.Document(io.BytesIO(content_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # Check tables
    table_texts = []
    has_tables = len(doc.tables) > 0
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)

    all_text = "\n".join(paragraphs + table_texts).strip()
    layout_meta = {
        "format": "docx",
        "hasColumns": False,
        "hasImages": False,
        "hasMultiColumnTables": has_tables,
        "contactInHeaderFooter": False,
        "hasNonStandardFonts": False,
        "pageCount": max(1, len(all_text.split()) // 350),
    }

    return all_text, layout_meta

def parse_resume_file(file_path: str, original_name: str) -> Dict[str, Any]:
    """Parse resume from disk into rawText and structured section data."""
    ext = os.path.splitext(original_name)[1].lower()
    if ext not in [".pdf", ".docx", ".txt"]:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")

    with open(file_path, "rb") as f:
        content_bytes = f.read()

    if len(content_bytes) == 0:
        raise ValueError("Uploaded file is empty (0 bytes).")

    validate_file_signature(content_bytes, ext)

    if ext == ".pdf":
        raw_text, layout_meta = parse_pdf_bytes(content_bytes)
    elif ext == ".docx":
        raw_text, layout_meta = parse_docx_bytes(content_bytes)
    else:  # .txt
        raw_text = content_bytes.decode("utf-8", errors="ignore")
        layout_meta = {
            "format": "txt",
            "hasColumns": False,
            "hasImages": False,
            "hasMultiColumnTables": False,
            "pageCount": 1
        }

    trimmed = raw_text.strip()
    word_count = len(trimmed.split())
    if word_count < 10:
        raise ValueError("Could not extract readable text from document. The file may be image-only, protected, or empty.")

    sections = extract_sections(trimmed)

    return {
        "rawText": trimmed,
        "structured": {
            "sections": sections,
            "layout": layout_meta,
            "wordCount": word_count,
            "lineCount": len(trimmed.split("\n")),
        }
    }
